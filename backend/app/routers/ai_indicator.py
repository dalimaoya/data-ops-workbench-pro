"""AI Indicator Design + Batch Create Tables — AI 指标设计 + 批量建表"""

import io
import json
import re
from typing import Optional, List

import httpx
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from pydantic import BaseModel

from app.database import get_db
from app.models import DatasourceConfig, AIConfig, UserAccount
from app.utils.auth import get_current_user, require_role
from app.utils.crypto import decrypt_password
from app.utils.remote_db import _connect
from app.utils.audit import log_operation
from app.i18n import t

router = APIRouter(prefix="/api/ai", tags=["ai-indicator"])


def _get_ai_config(db: Session):
    """Get AI configuration."""
    ai_cfg = db.query(AIConfig).first()
    if not ai_cfg:
        raise HTTPException(400, t("ai.config_required"))
    api_url = ai_cfg.cloud_api_url or ai_cfg.local_api_url or ai_cfg.api_url
    api_key = None
    model_name = ai_cfg.cloud_model_name or ai_cfg.local_model_name or ai_cfg.model_name
    api_protocol = ai_cfg.cloud_api_protocol or ai_cfg.local_api_protocol or ai_cfg.api_protocol or "openai"
    for key_field in [ai_cfg.cloud_api_key_encrypted, ai_cfg.local_api_key_encrypted, ai_cfg.api_key_encrypted]:
        if key_field:
            try:
                api_key = decrypt_password(key_field)
                break
            except Exception:
                continue
    if not api_url or not api_key:
        raise HTTPException(400, t("ai.config_invalid"))
    return api_url, api_key, model_name, api_protocol


async def _call_ai(api_url: str, api_key: str, model_name: str, api_protocol: str, prompt: str) -> str:
    """Call AI and return text response."""
    headers = {"Content-Type": "application/json"}
    if api_protocol == "claude":
        headers["x-api-key"] = api_key
        headers["anthropic-version"] = "2023-06-01"
        payload = {"model": model_name, "max_tokens": 4000, "messages": [{"role": "user", "content": prompt}]}
    else:
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {"model": model_name, "messages": [{"role": "user", "content": prompt}], "temperature": 0.3, "max_tokens": 4000}

    url = api_url.rstrip("/")
    if api_protocol == "claude":
        if not url.endswith("/messages"):
            url += "/v1/messages"
    else:
        if not url.endswith("/chat/completions"):
            if not url.endswith("/v1"):
                url += "/v1"
            url += "/chat/completions"

    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(url, json=payload, headers=headers)
        if resp.status_code != 200:
            raise HTTPException(502, t("ai.call_failed", detail=resp.text[:200]))
        result = resp.json()

    if api_protocol == "claude":
        return result.get("content", [{}])[0].get("text", "")
    return result.get("choices", [{}])[0].get("message", {}).get("content", "")


def _pinyin_initials(text: str) -> str:
    """Get simple pinyin initials for Chinese text (basic mapping)."""
    # Simple first-char approach for common characters
    result = ""
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff':
            result += ch
        elif ch.isalpha():
            result += ch.lower()
    # Truncate
    return result[:8] if result else "tbl"


@router.post("/design-indicators")
async def design_indicators(
    file: UploadFile = File(...),
    prefix: str = Form("ioc"),
    db: Session = Depends(get_db),
    user: UserAccount = Depends(require_role("admin")),
):
    """Upload a report file (Word/PDF/Excel) and AI parses indicator system."""
    content_text = ""
    filename = file.filename or ""

    data = await file.read()

    if filename.endswith(".xlsx") or filename.endswith(".xls"):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(max_row=200, values_only=True):
                content_text += " | ".join(str(c) if c else "" for c in row) + "\n"
    elif filename.endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(data))
        for para in doc.paragraphs[:300]:
            content_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                content_text += " | ".join(c.text for c in row.cells) + "\n"
    elif filename.endswith(".pdf"):
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages[:30]:
                text = page.extract_text()
                if text:
                    content_text += text + "\n"
                for table in (page.extract_tables() or []):
                    for row in table:
                        content_text += " | ".join(str(c) if c else "" for c in row) + "\n"
    else:
        content_text = data.decode("utf-8", errors="replace")[:10000]

    if not content_text.strip():
        raise HTTPException(400, t("ai.file_no_content"))

    # Truncate
    content_text = content_text[:8000]

    api_url, api_key, model_name, api_protocol = _get_ai_config(db)

    prompt = f"""你是一个数据指标设计专家。请分析以下报表内容，提取指标体系。

报表内容：
{content_text}

请返回一个 JSON 对象，格式如下：
{{
  "indicators": [
    {{
      "module": "模块名称",
      "name": "指标名称",
      "table_name": "建议表名（使用英文，下划线分隔，前缀为 {prefix}）",
      "fields": [
        {{"name": "字段名(英文)", "type": "数据类型(VARCHAR/INTEGER/DECIMAL等)", "comment": "中文注释"}}
      ]
    }}
  ]
}}

要求：
1. 表名规则：{prefix}_模块缩写_指标拼音首字母（如 {prefix}_szjj_czjr 表示数字经济-财政收入）
2. 每个指标对应一张表
3. 每张表都要有 id 主键字段
4. 数值字段用 DECIMAL(18,2)，文本字段用 VARCHAR
5. 只返回 JSON，不要其他内容"""

    ai_text = await _call_ai(api_url, api_key, model_name, api_protocol, prompt)

    # Parse JSON
    try:
        json_match = re.search(r'\{[\s\S]*\}', ai_text)
        if json_match:
            result = json.loads(json_match.group())
        else:
            result = {"_raw": ai_text}
    except json.JSONDecodeError:
        result = {"_raw": ai_text}

    log_operation(db, "AI指标设计", "解析指标", "success",
                  message=f"从文件 {filename} 解析出 {len(result.get('indicators', []))} 个指标",
                  operator=user.username)

    return result


class BatchCreateRequest(BaseModel):
    datasource_id: int
    indicators: list  # Same format as AI output
    schema_name: Optional[str] = None


@router.post("/batch-create-tables")
def batch_create_tables(
    body: BatchCreateRequest,
    db: Session = Depends(get_db),
    user: UserAccount = Depends(require_role("admin")),
):
    """Batch create tables from indicator design result."""
    ds = db.query(DatasourceConfig).filter(
        DatasourceConfig.id == body.datasource_id, DatasourceConfig.is_deleted == 0
    ).first()
    if not ds:
        raise HTTPException(404, t("ai_indicator.datasource_not_found"))

    pwd = decrypt_password(ds.password_encrypted)
    conn = _connect(ds.db_type, ds.host, ds.port, ds.username, pwd,
                    ds.database_name, body.schema_name, ds.charset, ds.connect_timeout_seconds or 10)

    created = []
    errors = []
    try:
        cur = conn.cursor()
        for indicator in body.indicators:
            table_name = indicator.get("table_name", "")
            fields_def = indicator.get("fields", [])
            if not table_name or not fields_def:
                continue

            cols_sql = []
            for fld in fields_def:
                fname = fld.get("name", "")
                ftype = fld.get("type", "VARCHAR(255)")
                if not fname:
                    continue
                col_def = f"{fname} {ftype}"
                if fname.lower() == "id":
                    if ds.db_type == "mysql":
                        col_def = f"{fname} INT AUTO_INCREMENT PRIMARY KEY"
                    elif ds.db_type in ("postgresql", "kingbase"):
                        col_def = f"{fname} SERIAL PRIMARY KEY"
                    elif ds.db_type == "sqlite":
                        col_def = f"{fname} INTEGER PRIMARY KEY AUTOINCREMENT"
                    else:
                        col_def = f"{fname} INTEGER PRIMARY KEY"
                cols_sql.append(col_def)

            if not cols_sql:
                continue

            ddl = f"CREATE TABLE {table_name} ({', '.join(cols_sql)})"
            try:
                cur.execute(ddl)
                created.append({
                    "table_name": table_name,
                    "module": indicator.get("module", ""),
                    "name": indicator.get("name", ""),
                    "field_count": len(cols_sql),
                })
            except Exception as e:
                errors.append({"table_name": table_name, "error": str(e)})

        conn.commit()
    finally:
        conn.close()

    # Auto-onboard created tables to platform
    from app.models import TableConfig, FieldConfig
    for tbl in created:
        # Check if already onboarded
        existing = db.query(TableConfig).filter(
            TableConfig.datasource_id == body.datasource_id,
            TableConfig.table_name == tbl["table_name"],
            TableConfig.is_deleted == 0,
        ).first()
        if existing:
            continue

        import uuid
        tc = TableConfig(
            table_config_code=f"TC-{uuid.uuid4().hex[:8].upper()}",
            datasource_id=body.datasource_id,
            db_name=ds.database_name,
            schema_name=body.schema_name,
            table_name=tbl["table_name"],
            table_alias=tbl.get("name", tbl["table_name"]),
            table_comment=f"{tbl.get('module', '')} - {tbl.get('name', '')}",
            primary_key_fields="id",
            created_by=user.username,
            updated_by=user.username,
        )
        db.add(tc)
        db.flush()

        # Add field configs
        indicator = next((i for i in body.indicators if i.get("table_name") == tbl["table_name"]), None)
        if indicator:
            for idx, fld in enumerate(indicator.get("fields", [])):
                fc = FieldConfig(
                    table_config_id=tc.id,
                    field_name=fld.get("name", ""),
                    display_name=fld.get("comment", fld.get("name", "")),
                    field_type=fld.get("type", "VARCHAR"),
                    display_order=idx + 1,
                    is_displayed=1,
                    is_editable=1 if fld.get("name", "").lower() != "id" else 0,
                    is_required=1 if fld.get("name", "").lower() == "id" else 0,
                    created_by=user.username,
                    updated_by=user.username,
                )
                db.add(fc)

    db.commit()

    log_operation(db, "AI指标设计", "批量建表", "success",
                  message=f"创建 {len(created)} 张表, 失败 {len(errors)} 张, 已自动纳管",
                  operator=user.username)

    return {"created": created, "errors": errors, "auto_managed": len(created)}
